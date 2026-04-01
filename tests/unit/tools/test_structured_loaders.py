import os
import pytest
import pandas as pd
from langchain_core.documents import Document
from tilellm.tools.structured_loaders import StructuredDocxLoader, ExcelLoader, CSVLoader

def test_csv_loader(tmp_path):
    csv_file = tmp_path / "test.csv"
    df = pd.DataFrame({
        "name": ["Alice", "Bob"],
        "age": [25, 30]
    })
    df.to_csv(csv_file, index=False)
    
    loader = CSVLoader(str(csv_file))
    docs = loader.load()
    
    assert len(docs) == 1
    assert "Alice" in docs[0].page_content
    assert "Bob" in docs[0].page_content
    assert docs[0].metadata["type"] == "csv"

def test_excel_loader(tmp_path):
    excel_file = tmp_path / "test.xlsx"
    df = pd.DataFrame({
        "name": ["Alice", "Bob"],
        "age": [25, 30]
    })
    df.to_excel(excel_file, index=False)
    
    loader = ExcelLoader(str(excel_file))
    docs = loader.load()
    
    assert len(docs) == 1
    assert "Alice" in docs[0].page_content
    assert docs[0].metadata["type"] == "xlsx"
    assert docs[0].metadata["sheet_name"] == "Sheet1"

def test_docx_loader(tmp_path):
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")
        
    docx_file = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_heading('Title 1', level=1)
    doc.add_paragraph('Hello world')
    doc.add_heading('Title 2', level=2)
    doc.add_paragraph('Content 2')
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Row 1 Col 1'
    table.cell(1, 1).text = 'Row 1 Col 2'
    
    doc.save(docx_file)
    
    loader = StructuredDocxLoader(str(docx_file))
    docs = loader.load()
    
    # Heading 1, Hello world, Heading 2, Content 2, Table
    assert len(docs) == 5
    
    assert docs[0].page_content == "Title 1"
    assert docs[0].metadata["heading_path"] == "Title 1"
    
    assert docs[1].page_content == "Hello world"
    assert docs[1].metadata["heading_path"] == "Title 1"
    
    assert docs[2].page_content == "Title 2"
    assert docs[2].metadata["heading_path"] == "Title 1 > Title 2"
    
    assert docs[3].page_content == "Content 2"
    assert docs[3].metadata["heading_path"] == "Title 1 > Title 2"
    
    assert "Header 1" in docs[4].page_content
    assert docs[4].metadata["element_type"] == "table"

def test_csv_loader_chunking(tmp_path):
    csv_file = tmp_path / "large.csv"
    df = pd.DataFrame({
        "id": range(150),
        "data": ["value"] * 150
    })
    df.to_csv(csv_file, index=False)
    
    loader = CSVLoader(str(csv_file))
    docs = loader.load()
    
    # 150 rows / 50 per chunk = 3 chunks
    assert len(docs) == 3
    assert docs[0].metadata["row_range"] == "0-50"
    assert docs[1].metadata["row_range"] == "50-100"
    assert docs[2].metadata["row_range"] == "100-150"
