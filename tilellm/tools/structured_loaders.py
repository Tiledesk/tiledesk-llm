import hashlib
import logging
import os
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
from langchain_core.documents import Document

try:
    import docx as python_docx
    from docx.oxml.ns import qn as docx_qn
except ImportError:
    python_docx = None
    docx_qn = None

logger = logging.getLogger(__name__)


class StructuredDocxLoader:
    """
    Load a DOCX file preserving heading hierarchy, tables, and (optionally)
    embedded images.

    Public API
    ----------
    load()                  → List[Document]          (text + tables, backward compat)
    load_with_images()      → (List[Document], List[image_record])
    extract_images(doc)     → List[image_record]       (call after _open_doc)

    image_record keys
    -----------------
    image_id      str   unique ID: docx_img_{para_index}_{md5[:8]}
    image_bytes   bytes raw blob from the DOCX package
    content_type  str   MIME type (image/png, image/jpeg, …)
    para_index    int   0-based paragraph index that contains the drawing
    alt_text      str   wp:docPr descr / name attribute (may be empty)
    """

    def __init__(self, file_path: str):
        self.file_path = file_path

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def load(self) -> List[Document]:
        """Return text + table Documents.  Backward-compatible."""
        doc, local_path, is_temp = self._open_doc()
        try:
            return self._extract_documents(doc)
        finally:
            if is_temp and os.path.exists(local_path):
                os.remove(local_path)

    def load_with_images(self) -> Tuple[List[Document], List[Dict[str, Any]]]:
        """
        Return (documents, image_records).
        Documents include text + tables with ``_para_index`` in metadata.
        image_records carry raw bytes + context needed for MinIO/LLM.
        """
        doc, local_path, is_temp = self._open_doc()
        try:
            documents = self._extract_documents(doc)
            image_records = self._extract_images(doc)
            return documents, image_records
        finally:
            if is_temp and os.path.exists(local_path):
                os.remove(local_path)

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _open_doc(self) -> Tuple[Any, str, bool]:
        """Download if URL, open python-docx Document.  Returns (doc, path, is_temp)."""
        if python_docx is None:
            raise ImportError(
                "python-docx is not installed. Install it with `pip install python-docx`."
            )
        local_path = self.file_path
        is_temp = False
        if self.file_path.startswith(("http://", "https://")):
            import tempfile
            import requests
            response = requests.get(self.file_path, timeout=60)
            response.raise_for_status()
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(response.content)
                local_path = tmp.name
                is_temp = True
        doc = python_docx.Document(local_path)
        return doc, local_path, is_temp

    def _extract_documents(self, doc) -> List[Document]:
        """
        Extract paragraphs and tables from an open python-docx Document.
        Paragraphs get ``_para_index`` in metadata for later cross-referencing
        with extracted images.
        """
        documents: List[Document] = []
        heading_stack: List[str] = []

        for para_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue

            if para.style and para.style.name.startswith("Heading"):
                try:
                    parts = para.style.name.split(" ")
                    if len(parts) >= 2 and parts[1].isdigit():
                        level = int(parts[1])
                        heading_stack = heading_stack[: level - 1] + [para.text.strip()]
                except Exception:
                    pass

            documents.append(Document(
                page_content=para.text.strip(),
                metadata={
                    "source": self.file_path,
                    "heading_path": " > ".join(heading_stack),
                    "style": para.style.name if para.style else "Normal",
                    "type": "docx",
                    "_para_index": para_idx,  # internal cross-ref key
                },
            ))

        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])

            if table_data:
                df = pd.DataFrame(table_data)
                table_text = df.to_markdown(index=False)
                documents.append(Document(
                    page_content=table_text,
                    metadata={
                        "source": self.file_path,
                        "heading_path": " > ".join(heading_stack),
                        "type": "docx",
                        "element_type": "table",
                        "table_index": i,
                    },
                ))

        return documents

    def _extract_images(self, doc) -> List[Dict[str, Any]]:
        """
        Scan all paragraphs for embedded images (modern w:drawing and
        legacy w:pict).  Returns one dict per unique image relationship.

        image_record keys: image_id, image_bytes, content_type, para_index, alt_text
        """
        if docx_qn is None:
            logger.warning("python-docx not available — skipping image extraction")
            return []

        images: List[Dict[str, Any]] = []
        seen_rids: set = set()

        for para_idx, para in enumerate(doc.paragraphs):
            for run in para.runs:

                # ── Modern DOCX: inline drawings ──────────────────────────
                for drawing in run.element.findall(".//" + docx_qn("w:drawing")):
                    for blip in drawing.findall(".//" + docx_qn("a:blip")):
                        r_embed = blip.get(docx_qn("r:embed"))
                        if not r_embed or r_embed in seen_rids:
                            continue
                        if r_embed not in doc.part.rels:
                            continue
                        rel = doc.part.rels[r_embed]
                        if "image" not in rel.reltype.lower():
                            continue
                        seen_rids.add(r_embed)

                        image_bytes: bytes = rel.target_part.blob
                        content_type: str = (
                            getattr(rel.target_part, "content_type", None)
                            or "image/png"
                        )
                        img_hash = hashlib.md5(image_bytes).hexdigest()[:8]
                        image_id = f"docx_img_{para_idx}_{img_hash}"

                        # Alt text from wp:docPr
                        alt_text = ""
                        for docPr in drawing.findall(".//" + docx_qn("wp:docPr")):
                            alt_text = (
                                docPr.get("descr", "")
                                or docPr.get("name", "")
                                or ""
                            )

                        images.append({
                            "image_id": image_id,
                            "image_bytes": image_bytes,
                            "content_type": content_type,
                            "para_index": para_idx,
                            "alt_text": alt_text.strip(),
                        })

                # ── Legacy DOCX: w:pict / v:imagedata ─────────────────────
                for pict in run.element.findall(".//" + docx_qn("w:pict")):
                    for imagedata in pict.findall(".//" + docx_qn("v:imagedata")):
                        r_id = imagedata.get(docx_qn("r:id"))
                        if not r_id or r_id in seen_rids:
                            continue
                        if r_id not in doc.part.rels:
                            continue
                        rel = doc.part.rels[r_id]
                        if "image" not in rel.reltype.lower():
                            continue
                        seen_rids.add(r_id)

                        image_bytes = rel.target_part.blob
                        content_type = (
                            getattr(rel.target_part, "content_type", None)
                            or "image/png"
                        )
                        img_hash = hashlib.md5(image_bytes).hexdigest()[:8]
                        image_id = f"docx_img_{para_idx}_{img_hash}"

                        images.append({
                            "image_id": image_id,
                            "image_bytes": image_bytes,
                            "content_type": content_type,
                            "para_index": para_idx,
                            "alt_text": "",
                        })

        logger.info("Extracted %d embedded image(s) from DOCX (%s)", len(images), self.file_path)
        return images


# ---------------------------------------------------------------------------
# ExcelLoader
# ---------------------------------------------------------------------------

# Righe iniziali entro cui cercare l'intestazione reale (i modelli PA ne premettono
# una o due di istruzioni; oltre questa soglia si tratta di un foglio anomalo).
_HEADER_SCAN_ROWS = 10

# Oltre questa dimensione il markdown della tabella viene spezzato per righe.
# Il taglio deve avvenire QUI e non nello splitter a valle: `to_markdown` riscrive
# l'intestazione in ogni blocco, mentre un taglio a caratteri la lascerebbe solo
# nel primo chunk, rendendo gli altri illeggibili (celle senza nome di colonna).
_MAX_TABLE_CHARS = 4000
_ROWS_PER_CHUNK = 50


def _detect_header_row(path, sheet_name) -> int:
    """Indice (0-based) della riga di intestazione reale del foglio.

    I modelli della PA premettono righe di nota ("L'OE e' tenuto ad aggiungere una
    riga per ogni ref. offerto..."), e la loro quantita' cambia da allegato ad
    allegato. Assumendo la riga 0 si prendono quelle note come intestazione e ogni
    colonna vera diventa `Unnamed: N`. Si sceglie invece la prima riga con il numero
    massimo di celle piene fra le prime righe del foglio.
    """
    try:
        probe = pd.read_excel(path, sheet_name=sheet_name, header=None,
                              nrows=_HEADER_SCAN_ROWS)
    except Exception as e:
        logger.debug("Header detection non riuscita su '%s' (%s): uso la riga 0.", sheet_name, e)
        return 0
    if probe.empty:
        return 0
    filled = probe.notna().sum(axis=1)
    best = int(filled.max())
    if best <= 1:
        # Nessuna riga tabellare nelle prime righe: foglio di sole note.
        return 0
    return int(filled.idxmax())


class ExcelLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        local_path = self.file_path
        is_temp = False
        if self.file_path.startswith(("http://", "https://")):
            import requests
            import tempfile
            response = requests.get(self.file_path, timeout=60)
            response.raise_for_status()
            suffix = ".xlsx" if ".xlsx" in self.file_path.lower() else ".xls"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(response.content)
                local_path = tmp.name
                is_temp = True

        try:
            documents: List[Document] = []
            excel_file = pd.ExcelFile(local_path)
            for sheet_name in excel_file.sheet_names:
                header_row = _detect_header_row(local_path, sheet_name)
                df = pd.read_excel(excel_file, sheet_name=sheet_name, header=header_row)
                if df.empty:
                    continue

                # Colonne interamente vuote: informazione zero, costo altissimo.
                # Nei tracciati di gara reali sono 59 su 61 (il modello e' ampio, il
                # compilatore ne usa una minima parte): renderle produce ~86% di
                # padding `nan`, che satura la finestra dell'embedder prima del dato.
                slim = df.dropna(axis=1, how="all")
                if not slim.empty and slim.shape[1] < df.shape[1]:
                    logger.info(
                        "ExcelLoader '%s': scartate %d colonne vuote su %d.",
                        sheet_name, df.shape[1] - slim.shape[1], df.shape[1],
                    )
                    df = slim

                col_names = list(df.columns.astype(str))
                file_type = "xlsx" if local_path.endswith(".xlsx") else "xls"
                rows_per_chunk = _ROWS_PER_CHUNK
                rendered = df.to_markdown(index=False)
                # Il numero di righe da solo non basta: una tabella larga supera i
                # 45.000 caratteri gia' con 29 righe.
                if len(df) > 100 or len(rendered) > _MAX_TABLE_CHARS:
                    rows_per_chunk = min(
                        rows_per_chunk,
                        max(1, int(len(df) * _MAX_TABLE_CHARS / max(len(rendered), 1))),
                    )
                    for i in range(0, len(df), rows_per_chunk):
                        chunk_df = df.iloc[i : i + rows_per_chunk]
                        documents.append(Document(
                            page_content=chunk_df.to_markdown(index=False),
                            metadata={
                                "source": self.file_path,
                                "sheet_name": sheet_name,
                                "col_names": ", ".join(col_names),
                                "row_range": f"{i}-{i + len(chunk_df)}",
                                "type": file_type,
                                "element_type": "table",
                            },
                        ))
                else:
                    documents.append(Document(
                        page_content=df.to_markdown(index=False),
                        metadata={
                            "source": self.file_path,
                            "sheet_name": sheet_name,
                            "col_names": ", ".join(col_names),
                            "type": file_type,
                            "element_type": "table",
                        },
                    ))
            return documents
        finally:
            if is_temp and os.path.exists(local_path):
                os.remove(local_path)


# ---------------------------------------------------------------------------
# CSVLoader
# ---------------------------------------------------------------------------

class CSVLoader:
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        local_path = self.file_path
        is_temp = False
        if self.file_path.startswith(("http://", "https://")):
            import requests
            import tempfile
            response = requests.get(self.file_path, timeout=60)
            response.raise_for_status()
            with tempfile.NamedTemporaryFile(delete=False, suffix=".csv") as tmp:
                tmp.write(response.content)
                local_path = tmp.name
                is_temp = True

        try:
            df = pd.read_csv(local_path)
            if df.empty:
                return []

            col_names = list(df.columns.astype(str))
            documents: List[Document] = []
            rows_per_chunk = 50
            if len(df) > 100:
                for i in range(0, len(df), rows_per_chunk):
                    chunk_df = df.iloc[i : i + rows_per_chunk]
                    documents.append(Document(
                        page_content=chunk_df.to_markdown(index=False),
                        metadata={
                            "source": self.file_path,
                            "col_names": ", ".join(col_names),
                            "row_range": f"{i}-{i + len(chunk_df)}",
                            "type": "csv",
                            "element_type": "table",
                        },
                    ))
            else:
                documents.append(Document(
                    page_content=df.to_markdown(index=False),
                    metadata={
                        "source": self.file_path,
                        "col_names": ", ".join(col_names),
                        "type": "csv",
                        "element_type": "table",
                    },
                ))
            return documents
        finally:
            if is_temp and os.path.exists(local_path):
                os.remove(local_path)
