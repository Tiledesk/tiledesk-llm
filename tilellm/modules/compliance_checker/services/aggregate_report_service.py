"""
AggregateReportService — combine multiple completed v2 checks into a
cross-operator comparison report.

Input  : N URLs of ComplianceReportV2 JSON documents (one per operator/check).
Output : an AggregateReport grouped by lot, rendered to md / docx / pdf / json.

Design:
  - Operator label = explicit per-source label, else derived from the namespace
    (convention gara-{anno}-lotto{N}-{operatore}).
  - Reports are grouped by lot_id (one comparison table per lot).
  - Per-operator synthesis is deterministic by default; an optional LLM synthesis
    is used only when the request enables it (and credentials are present).
  - docx (python-docx) and pdf (reportlab) renderers are guarded: if the library
    is unavailable the export falls back to Markdown with a warning.
"""
import io
import json
import logging
import re
from typing import List, Optional, Tuple

import httpx

from tilellm.modules.compliance_checker.models_v2 import (
    AggregateReport,
    AggregateReportRequest,
    ComplianceReportV2,
    HumanReviewItem,
    LotComparison,
    OperatorScore,
)
from tilellm.shared.utility import inject_llm_chat_async

logger = logging.getLogger(__name__)

MAX_REPORT_SIZE: int = 5 * 1024 * 1024  # 5 MB per report JSON

# gara-{anno}-lotto{N}-{operatore}  →  capture the operator segment
_OPERATOR_RE = re.compile(r"lotto\d+-(.+)$", re.IGNORECASE)

_LLM_SYNTHESIS_SYSTEM = (
    "Sei un assistente di una commissione di gara d'appalto pubblica. "
    "Sintetizza in 1-2 frasi in italiano la posizione di un operatore economico, "
    "basandoti SOLO sui dati forniti (punteggi e criteri da revisionare). "
    "Tono neutro e fattuale, nessuna raccomandazione di aggiudicazione."
)


class AggregateReportService:
    def __init__(self, llm=None):
        self._llm = llm

    # ------------------------------------------------------------------
    # Public pipeline
    # ------------------------------------------------------------------

    async def build(self, request: AggregateReportRequest) -> AggregateReport:
        """Download each report, group by lot, compute per-operator scores."""
        scores: List[OperatorScore] = []
        for source in request.sources:
            report = await self._download_report(source.url)
            label = source.operator_label or self._derive_operator(report.namespace)
            score = self._operator_score(report, label)
            if request.synthesis_llm and self._llm is not None:
                score.synthesis = await self._llm_synthesis(score)
            else:
                score.synthesis = self._deterministic_synthesis(score)
            scores.append(score)

        # Group by lot_id, sort operators by AI points desc
        lots: dict[str, LotComparison] = {}
        for sc in scores:
            lot = lots.get(sc.lot_id)
            if lot is None:
                lot = LotComparison(lot_id=sc.lot_id, lot_name=sc.lot_name, operators=[])
                lots[sc.lot_id] = lot
            lot.operators.append(sc)
        for lot in lots.values():
            lot.operators.sort(key=lambda o: o.ai_scored_points, reverse=True)

        return AggregateReport(
            lots=list(lots.values()),
            total_operators=len(scores),
            total_reports=len(request.sources),
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    async def _download_report(self, url: str) -> ComplianceReportV2:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(url)
            response.raise_for_status()
            content = response.content
            if len(content) > MAX_REPORT_SIZE:
                raise ValueError(
                    f"Il report è troppo grande ({len(content)} byte > {MAX_REPORT_SIZE} massimi)."
                )
        try:
            data = json.loads(content.decode("utf-8", errors="replace"))
        except json.JSONDecodeError as e:
            raise ValueError(f"Report non è JSON valido ({url}): {e}") from e
        return ComplianceReportV2.model_validate(data)

    @staticmethod
    def _derive_operator(namespace: str) -> str:
        m = _OPERATOR_RE.search(namespace or "")
        return m.group(1) if m else (namespace or "operatore")

    def _operator_score(self, report: ComplianceReportV2, label: str) -> OperatorScore:
        s = report.summary
        return OperatorScore(
            operator_label=label,
            namespace=report.namespace,
            lot_id=report.tender.lot_id,
            lot_name=report.tender.lot_name,
            ai_scored_points=s.ai_scored_points,
            ai_max_scorable_points=s.ai_max_scorable_points,
            tabular_total=s.tabular.total,
            tabular_compliant=s.tabular.compliant,
            tabular_compliance_rate=s.tabular.compliance_rate,
            human_review_count=s.human_review_count,
            citation_unattributed_count=s.citation_unattributed_count,
            human_review_items=self._human_review_items(report),
            tabular_results=report.tabular_results,
            discretionary_results=report.discretionary_results,
        )

    @staticmethod
    def _human_review_items(report: ComplianceReportV2) -> List[HumanReviewItem]:
        items: List[HumanReviewItem] = []
        # Discretionary criteria explicitly flagged
        for d in report.discretionary_results:
            if d.human_review_required:
                items.append(HumanReviewItem(
                    item_id=d.criterion_id,
                    item_type="discretionary",
                    reason=d.human_review_reason or "Revisione umana richiesta.",
                ))
        # Tabular requirements that could not be auto-verified
        for r in report.tabular_results:
            if r.judgment == "not_verifiable":
                items.append(HumanReviewItem(
                    item_id=r.requirement_id,
                    item_type="tabular",
                    reason=r.justification or "Requisito non verificabile automaticamente.",
                ))
        return items

    @staticmethod
    def _deterministic_synthesis(score: OperatorScore) -> str:
        pct = (
            round(100 * score.ai_scored_points / score.ai_max_scorable_points)
            if score.ai_max_scorable_points > 0 else 0
        )
        parts = [
            f"Punteggio IA {score.ai_scored_points:.2f}/{score.ai_max_scorable_points:.2f} ({pct}%)"
        ]
        if score.tabular_total > 0:
            parts.append(
                f"requisiti minimi {score.tabular_compliant}/{score.tabular_total} conformi "
                f"({score.tabular_compliance_rate:.0%})"
            )
        if score.human_review_count > 0:
            parts.append(f"{score.human_review_count} criteri in revisione umana")
        if score.citation_unattributed_count > 0:
            parts.append(f"{score.citation_unattributed_count} citazioni da verificare")
        return "; ".join(parts) + "."

    async def _llm_synthesis(self, score: OperatorScore) -> str:
        from langchain_core.messages import HumanMessage, SystemMessage
        facts = self._deterministic_synthesis(score)
        review = ", ".join(f"{it.item_id} ({it.reason})" for it in score.human_review_items) or "nessuno"
        user = (
            f"Operatore: {score.operator_label} — Lotto {score.lot_id} ({score.lot_name}).\n"
            f"Dati: {facts}\n"
            f"Criteri da revisionare: {review}.\n"
            "Scrivi una sintesi di 1-2 frasi."
        )
        try:
            resp = await self._llm.ainvoke([
                SystemMessage(content=_LLM_SYNTHESIS_SYSTEM),
                HumanMessage(content=user),
            ])
            content = resp.content
            if isinstance(content, list):
                content = " ".join(
                    p.get("text", "") if isinstance(p, dict) else str(p) for p in content
                )
            text = str(content).strip()
            return text or self._deterministic_synthesis(score)
        except Exception as e:
            logger.warning("LLM synthesis failed for '%s': %s — using deterministic.", score.operator_label, e)
            return self._deterministic_synthesis(score)


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------

def _pct(points: float, maximum: float) -> str:
    return f"{round(100 * points / maximum)}%" if maximum > 0 else "—"


def _discretionary_score_cell(d) -> str:
    if d.human_review_required:
        return "⚠ REVISIONE"
    if d.score is not None:
        return f"{d.score:.2f}"
    if d.measured_value:
        return f"Valore: {d.measured_value}"
    return "N/V"


def _citation_cell(d) -> str:
    """Fonte documentale che sostiene il punteggio: documento + pagina (+ sezione)."""
    if not d.citation_attributed:
        return "⚠ da verificare"
    if d.evidence_document:
        cell = f"{d.evidence_document} p.{d.evidence_page}"
        if d.evidence_section:
            cell += f" — {d.evidence_section}"
        return cell
    return "—"


def render_markdown(agg: AggregateReport) -> str:
    """Render the aggregate report as a Markdown document.

    Per lotto, tre blocchi nell'ordine: tabella di sintesi comparativa →
    dettaglio per ogni operatore → sintesi narrativa.
    """
    buf = io.StringIO()
    buf.write("# Report comparativo offerte\n\n")
    buf.write(f"**Operatori:** {agg.total_operators} | **Report combinati:** {agg.total_reports}\n\n")

    for lot in agg.lots:
        buf.write(f"## Lotto {lot.lot_id} — {lot.lot_name}\n\n")

        # --- Blocco 1: tabella di sintesi comparativa ---
        buf.write("### Tabella di sintesi\n\n")
        buf.write(
            "| # | Operatore | Punteggio IA | Max IA | % | Tabellari | Revisione umana | Citazioni da verificare |\n"
        )
        buf.write("|---|-----------|-------------|--------|---|-----------|-----------------|------------------------|\n")
        for i, op in enumerate(lot.operators, 1):
            tab = f"{op.tabular_compliant}/{op.tabular_total}" if op.tabular_total else "—"
            label = op.operator_label.replace("|", "\\|")
            buf.write(
                f"| {i} | {label} | {op.ai_scored_points:.2f} | {op.ai_max_scorable_points:.2f} "
                f"| {_pct(op.ai_scored_points, op.ai_max_scorable_points)} | {tab} "
                f"| {op.human_review_count} | {op.citation_unattributed_count} |\n"
            )
        buf.write("\n")

        # --- Blocco 2: dettaglio per ogni operatore ---
        buf.write("### Dettaglio per operatore\n\n")
        for op in lot.operators:
            buf.write(f"#### {op.operator_label}\n\n")
            if op.discretionary_results:
                buf.write("| ID | Criterio | Modalità | Punteggio | Max | Revisione | Fonte (doc/pag.) |\n")
                buf.write("|----|----------|----------|-----------|-----|-----------|------------------|\n")
                for d in op.discretionary_results:
                    text = d.criterion_text.replace("|", "\\|")
                    cit = _citation_cell(d).replace("|", "\\|")
                    rev = "sì" if d.human_review_required else "—"
                    buf.write(
                        f"| {d.criterion_id} | {text} | {d.mode.value} | {_discretionary_score_cell(d)} "
                        f"| {d.max_points:g} | {rev} | {cit} |\n"
                    )
                buf.write("\n")
            if op.tabular_results:
                buf.write("| Requisito | Esito | Documento | Pagina |\n")
                buf.write("|-----------|-------|-----------|--------|\n")
                for r in op.tabular_results:
                    req = r.requirement_text.replace("|", "\\|")
                    doc = (r.evidence_document or "").replace("|", "\\|")
                    page = str(r.evidence_page) if r.evidence_document else ""
                    buf.write(f"| {req} | {r.judgment} | {doc} | {page} |\n")
                buf.write("\n")
            if op.human_review_items:
                buf.write("**Punti da revisionare:**\n\n")
                for it in op.human_review_items:
                    buf.write(f"- `{it.item_id}` ({it.item_type}): {it.reason}\n")
                buf.write("\n")

        # --- Blocco 3: sintesi narrativa ---
        buf.write("### Sintesi\n\n")
        for op in lot.operators:
            buf.write(f"- **{op.operator_label}**: {op.synthesis}\n")
        buf.write("\n")

    return buf.getvalue()


def render_docx(agg: AggregateReport) -> bytes:
    """Render to DOCX. Raises ImportError if python-docx is unavailable."""
    from docx import Document  # guarded import

    doc = Document()
    doc.add_heading("Report comparativo offerte", level=0)
    doc.add_paragraph(f"Operatori: {agg.total_operators} | Report combinati: {agg.total_reports}")

    for lot in agg.lots:
        doc.add_heading(f"Lotto {lot.lot_id} — {lot.lot_name}", level=1)

        # Blocco 1: tabella di sintesi comparativa
        doc.add_heading("Tabella di sintesi", level=2)
        headers = ["#", "Operatore", "Punteggio IA", "Max IA", "%", "Tabellari", "Revisione umana", "Citazioni"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for c, h in enumerate(headers):
            table.rows[0].cells[c].text = h
        for i, op in enumerate(lot.operators, 1):
            tab = f"{op.tabular_compliant}/{op.tabular_total}" if op.tabular_total else "—"
            row = table.add_row().cells
            values = [
                str(i), op.operator_label, f"{op.ai_scored_points:.2f}",
                f"{op.ai_max_scorable_points:.2f}", _pct(op.ai_scored_points, op.ai_max_scorable_points), tab,
                str(op.human_review_count), str(op.citation_unattributed_count),
            ]
            for c, v in enumerate(values):
                row[c].text = v

        # Blocco 2: dettaglio per operatore
        doc.add_heading("Dettaglio per operatore", level=2)
        for op in lot.operators:
            doc.add_heading(op.operator_label, level=3)
            if op.discretionary_results:
                dheaders = ["ID", "Criterio", "Modalità", "Punteggio", "Max", "Revisione", "Fonte (doc/pag.)"]
                dt = doc.add_table(rows=1, cols=len(dheaders))
                dt.style = "Light Grid Accent 1"
                for c, h in enumerate(dheaders):
                    dt.rows[0].cells[c].text = h
                for d in op.discretionary_results:
                    cells = dt.add_row().cells
                    vals = [
                        d.criterion_id, d.criterion_text, d.mode.value,
                        _discretionary_score_cell(d), f"{d.max_points:g}",
                        "sì" if d.human_review_required else "—",
                        _citation_cell(d),
                    ]
                    for c, v in enumerate(vals):
                        cells[c].text = v
            if op.human_review_items:
                doc.add_paragraph("Punti da revisionare:")
                for it in op.human_review_items:
                    doc.add_paragraph(f"{it.item_id} ({it.item_type}): {it.reason}", style="List Bullet")

        # Blocco 3: sintesi narrativa
        doc.add_heading("Sintesi", level=2)
        for op in lot.operators:
            doc.add_paragraph(f"{op.operator_label}: {op.synthesis}", style="List Bullet")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render_pdf(agg: AggregateReport) -> bytes:
    """Render to PDF. Raises ImportError if reportlab is unavailable."""
    from reportlab.lib import colors  # guarded import
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    )

    styles = getSampleStyleSheet()
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=landscape(A4))
    flow = [
        Paragraph("Report comparativo offerte", styles["Title"]),
        Paragraph(f"Operatori: {agg.total_operators} | Report combinati: {agg.total_reports}", styles["Normal"]),
        Spacer(1, 12),
    ]
    header_style = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f2f2")]),
    ])

    for lot in agg.lots:
        flow.append(Paragraph(f"Lotto {lot.lot_id} — {lot.lot_name}", styles["Heading1"]))

        # Blocco 1: tabella di sintesi comparativa
        flow.append(Paragraph("Tabella di sintesi", styles["Heading2"]))
        data = [["#", "Operatore", "Punt. IA", "Max IA", "%", "Tabellari", "Rev. umana", "Citazioni"]]
        for i, op in enumerate(lot.operators, 1):
            tab = f"{op.tabular_compliant}/{op.tabular_total}" if op.tabular_total else "—"
            data.append([
                str(i), op.operator_label, f"{op.ai_scored_points:.2f}",
                f"{op.ai_max_scorable_points:.2f}", _pct(op.ai_scored_points, op.ai_max_scorable_points), tab,
                str(op.human_review_count), str(op.citation_unattributed_count),
            ])
        summary_tbl = Table(data, repeatRows=1)
        summary_tbl.setStyle(header_style)
        flow.append(summary_tbl)
        flow.append(Spacer(1, 8))

        # Blocco 2: dettaglio per operatore
        flow.append(Paragraph("Dettaglio per operatore", styles["Heading2"]))
        for op in lot.operators:
            flow.append(Paragraph(op.operator_label, styles["Heading3"]))
            if op.discretionary_results:
                ddata = [["ID", "Criterio", "Modalità", "Punteggio", "Max", "Revisione", "Fonte (doc/pag.)"]]
                for d in op.discretionary_results:
                    ddata.append([
                        d.criterion_id,
                        Paragraph(d.criterion_text, styles["Normal"]),
                        d.mode.value, _discretionary_score_cell(d), f"{d.max_points:g}",
                        "sì" if d.human_review_required else "—",
                        Paragraph(_citation_cell(d), styles["Normal"]),
                    ])
                dtbl = Table(ddata, repeatRows=1)
                dtbl.setStyle(header_style)
                flow.append(dtbl)
            for it in op.human_review_items:
                flow.append(Paragraph(f"• {it.item_id} ({it.item_type}): {it.reason}", styles["Normal"]))
            flow.append(Spacer(1, 6))

        # Blocco 3: sintesi narrativa
        flow.append(Paragraph("Sintesi", styles["Heading2"]))
        for op in lot.operators:
            flow.append(Paragraph(f"• {op.operator_label}: {op.synthesis}", styles["Normal"]))
        flow.append(Spacer(1, 8))

    doc.build(flow)
    return out.getvalue()


def render_report(agg: AggregateReport, output_format: str) -> Tuple[bytes, str, Optional[str]]:
    """
    Render to the requested format. Returns (content_bytes, actual_format, warning).

    docx/pdf fall back to Markdown (with a warning) when their library is missing.
    """
    fmt = output_format.lower()
    if fmt == "docx":
        try:
            return render_docx(agg), "docx", None
        except ImportError:
            logger.warning("python-docx not available — falling back to Markdown for docx export.")
            return render_markdown(agg).encode("utf-8"), "md", (
                "Formato 'docx' non disponibile (libreria assente): esportato come Markdown."
            )
    if fmt == "pdf":
        try:
            return render_pdf(agg), "pdf", None
        except ImportError:
            logger.warning("reportlab not available — falling back to Markdown for pdf export.")
            return render_markdown(agg).encode("utf-8"), "md", (
                "Formato 'pdf' non disponibile (libreria assente): esportato come Markdown."
            )
    # md (default)
    return render_markdown(agg).encode("utf-8"), "md", None


# ---------------------------------------------------------------------------
# DI entry point
# ---------------------------------------------------------------------------

@inject_llm_chat_async
async def _resolve_llm(request, llm=None, **kwargs):
    """Resolve a chat LLM from the request credentials (used only for synthesis)."""
    return llm


async def build_aggregate_report(request: AggregateReportRequest) -> AggregateReport:
    """
    Build the aggregate report. Resolves an LLM only when synthesis_llm is enabled,
    so the deterministic path needs no LLM credentials.
    """
    llm = None
    if request.synthesis_llm:
        try:
            llm = await _resolve_llm(request)
        except Exception as e:
            logger.warning("Could not resolve LLM for synthesis: %s — using deterministic synthesis.", e)
            llm = None
    svc = AggregateReportService(llm=llm)
    return await svc.build(request)
